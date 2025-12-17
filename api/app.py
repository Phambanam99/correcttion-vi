# -*- coding: utf-8 -*-
"""
Flask API cho Vietnamese Text Corrector
Supports multiple models and pipeline strategies:
- qwen_protonx: Qwen (local) + ProtonX
- qwen_only: Qwen only (local)
- protonx_only: ProtonX only
- bartpho_protonx: BartPho + ProtonX
- ollama_protonx: Ollama (online) + ProtonX
- ollama_only: Ollama only (online)
"""

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import sys
import os
import queue
import threading
import uuid
from datetime import datetime, timedelta

# Add parent directory to path
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from config import QWEN_MODELS, PIPELINE_STRATEGIES, DEFAULT_PIPELINE, MAX_QUEUE_SIZE, JOB_TIMEOUT_SECONDS, JOB_CLEANUP_HOURS
from llm.bartpho_model import correct_text as bartpho_correct, correct_text_chunked as bartpho_chunked
from llm.qwen_model import correct_text as qwen_correct, get_available_models as get_qwen_models
from protonx_layer.protonx_refine import refine_text_chunked
from processor.diff_utils import generate_change_note, is_meaningful_text

# Load Ollama model
ollama_models_list = []
try:
    from llm.ollama_model import correct_text as ollama_correct, check_ollama_health, get_available_models as get_ollama_models
    ollama_available = check_ollama_health()
    if ollama_available:
        print("✅ Ollama API is reachable")
        ollama_models_list = get_ollama_models()
    else:
        print("⚠️ Ollama API không khả dụng")
except Exception as e:
    print(f"⚠️ Ollama module error: {e}")
    ollama_available = False
    ollama_correct = None

# Load Vistral model (gated model, cần HF_TOKEN)
vistral_available = False
vistral_correct = None

try:
    from llm.vistral_model import correct_text as vistral_correct
    vistral_available = True
    print("✅ Vistral model loaded successfully")
except Exception as e:
    print(f"⚠️ Vistral model không khả dụng: {e}")
    vistral_available = False

app = Flask(__name__)
CORS(app)  # Enable CORS for web frontend

# Cấu hình
MAX_WORDS_PER_CHUNK = 100

# Available models: base models + qwen variants (ollama models are fetched dynamically)
AVAILABLE_MODELS = ["bartpho", "qwen", "vistral"] + [f"qwen-{k}" for k in QWEN_MODELS.keys()]
DEFAULT_MODEL = "qwen"

# ===== JOB QUEUE SYSTEM =====
# Job statuses
JOB_STATUS_PENDING = "pending"
JOB_STATUS_PROCESSING = "processing"
JOB_STATUS_COMPLETED = "completed"
JOB_STATUS_FAILED = "failed"

# In-memory job store
job_store = {}  # {job_id: {status, created_at, result, error, ...}}
job_queue = queue.Queue(maxsize=MAX_QUEUE_SIZE)
job_store_lock = threading.Lock()


def job_worker():
    """Background worker thread to process jobs from queue"""
    while True:
        try:
            job_id = job_queue.get(timeout=1)
        except queue.Empty:
            continue
        
        try:
            with job_store_lock:
                if job_id not in job_store:
                    continue
                job = job_store[job_id]
                job["status"] = JOB_STATUS_PROCESSING
                job["started_at"] = datetime.now().isoformat()
            
            # Process the job
            text = job["text"]
            pipeline = job.get("pipeline", DEFAULT_PIPELINE)
            qwen_variant = job.get("qwen_model")
            ollama_model = job.get("ollama_model")
            
            # Execute correction
            final_text, explanation = correct_with_pipeline(
                text, 
                pipeline=pipeline, 
                qwen_variant=qwen_variant, 
                ollama_model=ollama_model
            )
            
            note = generate_change_note(text, final_text)
            
            with job_store_lock:
                job_store[job_id].update({
                    "status": JOB_STATUS_COMPLETED,
                    "completed_at": datetime.now().isoformat(),
                    "result": {
                        "original": text,
                        "corrected": final_text,
                        "explanation": explanation,
                        "note": note or "",
                        "has_changes": text != final_text
                    }
                })
            
            print(f"✅ Job {job_id[:8]}... completed")
            
        except Exception as e:
            import traceback
            with job_store_lock:
                if job_id in job_store:
                    job_store[job_id].update({
                        "status": JOB_STATUS_FAILED,
                        "completed_at": datetime.now().isoformat(),
                        "error": str(e),
                        "traceback": traceback.format_exc()
                    })
            print(f"❌ Job {job_id[:8]}... failed: {e}")
        
        finally:
            job_queue.task_done()


def cleanup_old_jobs():
    """Remove completed jobs older than JOB_CLEANUP_HOURS"""
    cutoff = datetime.now() - timedelta(hours=JOB_CLEANUP_HOURS)
    with job_store_lock:
        to_remove = []
        for job_id, job in job_store.items():
            if job["status"] in [JOB_STATUS_COMPLETED, JOB_STATUS_FAILED]:
                created = datetime.fromisoformat(job["created_at"])
                if created < cutoff:
                    to_remove.append(job_id)
        for job_id in to_remove:
            del job_store[job_id]
        if to_remove:
            print(f"🧹 Cleaned up {len(to_remove)} old jobs")


# Start worker thread
worker_thread = threading.Thread(target=job_worker, daemon=True)
worker_thread.start()
print("🔄 Job worker thread started")


def correct_with_model(text: str, model: str = DEFAULT_MODEL, qwen_variant: str = None) -> tuple:
    """
    Sửa lỗi văn bản với model được chọn.
    Returns: (corrected_text, explanation)
    
    Args:
        text: Văn bản cần sửa
        model: Model chính (bartpho, qwen, vistral, hoặc qwen-<variant>)
        qwen_variant: Variant của Qwen model (qwen2.5-7b, qwen3-8b)
    """
    word_count = len(text.split())
    
    # Handle qwen-<variant> format
    if model.startswith("qwen-"):
        qwen_variant = model.replace("qwen-", "")
        model = "qwen"
    
    if model == "qwen":
        # Qwen trả về tuple (text, explanation)
        corrected, explanation = qwen_correct(text, model_key=qwen_variant)
        return corrected, explanation
    elif model == "vistral":
        # Vistral model
        if vistral_available and vistral_correct:
            corrected, explanation = vistral_correct(text)
            return corrected, explanation
        else:
            # Fallback to Qwen nếu Vistral không available
            print("⚠️ Vistral không khả dụng, dùng Qwen thay thế")
            corrected, explanation = qwen_correct(text)
            explanation = "⚠️ Vistral không khả dụng (cần HF_TOKEN). Đã dùng Qwen."
            return corrected, explanation
    else:
        # BartPho (default)
        if word_count > MAX_WORDS_PER_CHUNK:
            corrected = bartpho_chunked(text, MAX_WORDS_PER_CHUNK)
        else:
            corrected = bartpho_correct(text)
        explanation = generate_explanation(text, corrected)
        return corrected, explanation


def correct_with_pipeline(text: str, model: str = DEFAULT_MODEL, pipeline: str = DEFAULT_PIPELINE, qwen_variant: str = None, ollama_model: str = None) -> tuple:
    """
    Sửa lỗi văn bản với pipeline được chọn.
    Returns: (corrected_text, explanation)
    
    Pipeline strategies:
    - qwen_protonx: Qwen → ProtonX refine
    - qwen_only: Chỉ Qwen
    - protonx_only: Chỉ ProtonX
    - bartpho_protonx: BartPho → ProtonX refine
    - ollama_protonx: Ollama → ProtonX refine (online)
    - ollama_only: Chỉ Ollama (online)
    """
    word_count = len(text.split())
    
    if pipeline == "qwen_only":
        # Chỉ dùng Qwen, không ProtonX
        corrected, explanation = qwen_correct(text, model_key=qwen_variant)
        return corrected, explanation
    
    elif pipeline == "protonx_only":
        # Chỉ dùng ProtonX
        corrected = refine_text_chunked(text, MAX_WORDS_PER_CHUNK)
        explanation = "Đã refine với ProtonX (không qua LLM)"
        return corrected, explanation
    
    elif pipeline == "bartpho_protonx":
        # BartPho + ProtonX
        if word_count > MAX_WORDS_PER_CHUNK:
            model_fixed = bartpho_chunked(text, MAX_WORDS_PER_CHUNK)
        else:
            model_fixed = bartpho_correct(text)
        # ProtonX refine
        final_text = refine_text_chunked(model_fixed, MAX_WORDS_PER_CHUNK)
        explanation = generate_explanation(text, final_text)
        return final_text, explanation
    
    elif pipeline == "ollama_only":
        # Chỉ dùng Ollama (online), không ProtonX
        if ollama_available and ollama_correct:
            corrected, explanation = ollama_correct(text, model_key=ollama_model)
            return corrected, explanation
        else:
            # Fallback to Qwen
            print("⚠️ Ollama không khả dụng, dùng Qwen thay thế")
            corrected, explanation = qwen_correct(text, model_key=qwen_variant)
            explanation = "⚠️ Ollama API không khả dụng. Đã dùng Qwen local."
            return corrected, explanation
    
    elif pipeline == "ollama_protonx":
        # Ollama (online) + ProtonX
        if ollama_available and ollama_correct:
            model_fixed, explanation = ollama_correct(text, model_key=ollama_model)
        else:
            print("⚠️ Ollama không khả dụng, dùng Qwen thay thế")
            model_fixed, explanation = qwen_correct(text, model_key=qwen_variant)
            explanation = "⚠️ Ollama API không khả dụng. Đã dùng Qwen local."
        # ProtonX refine
        final_text = refine_text_chunked(model_fixed, MAX_WORDS_PER_CHUNK)
        return final_text, explanation
    
    else:  # qwen_protonx (default)
        # Qwen + ProtonX
        model_fixed, explanation = qwen_correct(text, model_key=qwen_variant)
        # ProtonX refine
        final_text = refine_text_chunked(model_fixed, MAX_WORDS_PER_CHUNK)
        return final_text, explanation


def generate_explanation(original: str, corrected: str) -> str:
    """Tạo giải thích ngắn gọn về các thay đổi"""
    if original.strip() == corrected.strip():
        return "Không có thay đổi."
    
    original_words = set(original.lower().split())
    corrected_words = set(corrected.lower().split())
    
    added = corrected_words - original_words
    removed = original_words - corrected_words
    
    explanations = []
    if removed:
        explanations.append(f"Sửa: {', '.join(list(removed)[:5])}")
    if added:
        explanations.append(f"Thành: {', '.join(list(added)[:5])}")
    
    return " → ".join(explanations) if explanations else "Đã sửa dấu và định dạng."


@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        "status": "ok",
        "message": "Vietnamese Text Corrector API is running",
        "available_models": AVAILABLE_MODELS,
        "qwen_models": list(QWEN_MODELS.keys()),
        "ollama_models": ollama_models_list,
        "ollama_available": ollama_available,
        "available_pipelines": PIPELINE_STRATEGIES,
        "default_model": DEFAULT_MODEL,
        "default_pipeline": DEFAULT_PIPELINE
    })


@app.route('/api/ollama-models', methods=['GET'])
def get_ollama_models_endpoint():
    """
    Get available Ollama models from the remote API
    
    Response:
    {
        "success": true,
        "available": true/false,
        "models": ["model1", "model2", ...]
    }
    """
    global ollama_models_list
    
    if not ollama_available:
        return jsonify({
            "success": True,
            "available": False,
            "models": [],
            "message": "Ollama API không khả dụng"
        })
    
    try:
        # Refresh models from API
        models = get_ollama_models()
        ollama_models_list = models
        
        return jsonify({
            "success": True,
            "available": True,
            "models": models
        })
    except Exception as e:
        return jsonify({
            "success": False,
            "available": False,
            "models": ollama_models_list,
            "error": str(e)
        })


@app.route('/api/correct', methods=['POST'])
def correct_text():
    """
    Sửa lỗi văn bản
    
    Request body:
    {
        "text": "văn bản cần sửa",
        "pipeline": "qwen_protonx" (optional)
    }
    """
    try:
        data = request.get_json()
        if not data or 'text' not in data:
            return jsonify({
                "success": False,
                "error": "Missing 'text' field in request body"
            }), 400
        
        original = data['text'].strip()
        if not original:
            return jsonify({
                "success": False,
                "error": "Text cannot be empty"
            }), 400
        
        # Lấy pipeline
        pipeline = data.get('pipeline', DEFAULT_PIPELINE)
        if pipeline not in PIPELINE_STRATEGIES:
            pipeline = DEFAULT_PIPELINE
        
        qwen_variant = data.get('qwen_model', None)
        
        # Sửa lỗi với pipeline
        final_text, explanation = correct_with_pipeline(original, pipeline=pipeline, qwen_variant=qwen_variant)
        
        # Tạo ghi chú thay đổi
        note = generate_change_note(original, final_text)
        
        return jsonify({
            "success": True,
            "original": original,
            "corrected": final_text,
            "explanation": explanation,
            "pipeline_used": pipeline,
            "note": note or ""
        })
        
    except Exception as e:
        import traceback
        return jsonify({
            "success": False,
            "error": str(e),
            "traceback": traceback.format_exc()
        }), 500


@app.route('/api/submit-job', methods=['POST'])
def submit_job():
    """
    Submit a text correction job to the queue (async processing).
    Returns immediately with a job ID that can be polled for status.
    
    Request body:
    {
        "text": "văn bản cần sửa",
        "pipeline": "qwen_protonx" (optional),
        "qwen_model": "qwen3-8b" (optional),
        "ollama_model": "qwen2.5:7b" (optional)
    }
    
    Response:
    {
        "success": true,
        "job_id": "uuid-string",
        "queue_position": 5,
        "message": "Job submitted successfully"
    }
    """
    try:
        data = request.get_json()
        if not data or 'text' not in data:
            return jsonify({
                "success": False,
                "error": "Missing 'text' field in request body"
            }), 400
        
        text = data['text'].strip()
        if not text:
            return jsonify({
                "success": False,
                "error": "Text cannot be empty"
            }), 400
        
        # Check if queue is full
        if job_queue.full():
            return jsonify({
                "success": False,
                "error": "Queue is full. Please try again later.",
                "queue_size": job_queue.qsize()
            }), 503
        
        # Create job
        job_id = str(uuid.uuid4())
        pipeline = data.get('pipeline', DEFAULT_PIPELINE)
        if pipeline not in PIPELINE_STRATEGIES:
            pipeline = DEFAULT_PIPELINE
        
        job = {
            "job_id": job_id,
            "text": text,
            "pipeline": pipeline,
            "qwen_model": data.get('qwen_model'),
            "ollama_model": data.get('ollama_model'),
            "status": JOB_STATUS_PENDING,
            "created_at": datetime.now().isoformat(),
            "result": None,
            "error": None
        }
        
        with job_store_lock:
            job_store[job_id] = job
        
        job_queue.put(job_id)
        
        # Cleanup old jobs periodically
        if len(job_store) > MAX_QUEUE_SIZE * 2:
            cleanup_old_jobs()
        
        return jsonify({
            "success": True,
            "job_id": job_id,
            "queue_position": job_queue.qsize(),
            "message": "Job submitted successfully"
        })
        
    except Exception as e:
        import traceback
        return jsonify({
            "success": False,
            "error": str(e),
            "traceback": traceback.format_exc()
        }), 500


@app.route('/api/job-status/<job_id>', methods=['GET'])
def get_job_status(job_id):
    """
    Get the status of a submitted job.
    
    Response when pending:
    {
        "success": true,
        "status": "pending",
        "queue_position": 3
    }
    
    Response when completed:
    {
        "success": true,
        "status": "completed",
        "result": {
            "original": "...",
            "corrected": "...",
            "explanation": "...",
            "has_changes": true
        }
    }
    """
    with job_store_lock:
        if job_id not in job_store:
            return jsonify({
                "success": False,
                "error": "Job not found"
            }), 404
        
        job = job_store[job_id].copy()
    
    response = {
        "success": True,
        "job_id": job_id,
        "status": job["status"],
        "created_at": job["created_at"]
    }
    
    if job["status"] == JOB_STATUS_PENDING:
        response["queue_position"] = job_queue.qsize()
    elif job["status"] == JOB_STATUS_PROCESSING:
        response["started_at"] = job.get("started_at")
    elif job["status"] == JOB_STATUS_COMPLETED:
        response["completed_at"] = job.get("completed_at")
        response["result"] = job.get("result")
    elif job["status"] == JOB_STATUS_FAILED:
        response["completed_at"] = job.get("completed_at")
        response["error"] = job.get("error")
    
    return jsonify(response)


@app.route('/api/queue-status', methods=['GET'])
def get_queue_status():
    """
    Get the current queue status.
    
    Response:
    {
        "success": true,
        "queue_size": 5,
        "max_queue_size": 50,
        "pending_jobs": 3,
        "processing_jobs": 1,
        "completed_jobs": 10
    }
    """
    with job_store_lock:
        pending = sum(1 for j in job_store.values() if j["status"] == JOB_STATUS_PENDING)
        processing = sum(1 for j in job_store.values() if j["status"] == JOB_STATUS_PROCESSING)
        completed = sum(1 for j in job_store.values() if j["status"] == JOB_STATUS_COMPLETED)
        failed = sum(1 for j in job_store.values() if j["status"] == JOB_STATUS_FAILED)
    
    return jsonify({
        "success": True,
        "queue_size": job_queue.qsize(),
        "max_queue_size": MAX_QUEUE_SIZE,
        "pending_jobs": pending,
        "processing_jobs": processing,
        "completed_jobs": completed,
        "failed_jobs": failed,
        "total_jobs": len(job_store)
    })


@app.route('/api/correct-paragraphs', methods=['POST'])
def correct_paragraphs():
    """
    Sửa lỗi nhiều đoạn văn (tách bằng newline)
    
    Request body:
    {
        "text": "đoạn 1\nđoạn 2\nđoạn 3",
        "model": "qwen" hoặc "bartpho" (mặc định: qwen),
        "pipeline": "qwen_protonx" hoặc "qwen_only" hoặc "protonx_only" hoặc "bartpho_protonx",
        "qwen_model": "qwen2.5-7b" hoặc "qwen3-8b" (optional)
    }
    """
    try:
        data = request.get_json()
        if not data or 'text' not in data:
            return jsonify({
                "success": False,
                "error": "Missing 'text' field in request body"
            }), 400
        
        text = data['text'].strip()
        if not text:
            return jsonify({
                "success": False,
                "error": "Text cannot be empty"
            }), 400
        
        # Lấy model và pipeline
        model = data.get('model', DEFAULT_MODEL).lower()
        pipeline = data.get('pipeline', DEFAULT_PIPELINE)
        qwen_variant = data.get('qwen_model', None)
        ollama_model_name = None
        
        # Handle ollama-<model> format
        if model.startswith("ollama-"):
            ollama_model_name = model.replace("ollama-", "")
            # Auto-switch to ollama pipeline if using ollama model
            if pipeline not in ["ollama_only", "ollama_protonx"]:
                pipeline = "ollama_protonx"  # Default to ollama + protonx
            model = "ollama"
        
        # Handle qwen-<variant> format
        elif model.startswith("qwen-"):
            qwen_variant = model.replace("qwen-", "")
            model = "qwen"
        
        # Validate pipeline
        if pipeline not in PIPELINE_STRATEGIES:
            pipeline = DEFAULT_PIPELINE
        
        # Chia thành các đoạn
        paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
        
        results = []
        corrected_paragraphs = []
        
        for i, original in enumerate(paragraphs):
            # Kiểm tra đoạn văn có ý nghĩa để xử lý hay không
            if not is_meaningful_text(original):
                # Bỏ qua đoạn không có ý nghĩa, giữ nguyên
                results.append({
                    "index": i,
                    "original": original,
                    "corrected": original,
                    "explanation": "Đoạn văn không có nội dung ý nghĩa để xử lý",
                    "note": "",
                    "has_changes": False,
                    "skipped": True
                })
                corrected_paragraphs.append(original)
                continue
            
            # Sửa lỗi với pipeline
            final_text, explanation = correct_with_pipeline(original, model=model, pipeline=pipeline, qwen_variant=qwen_variant, ollama_model=ollama_model_name)
            
            note = generate_change_note(original, final_text)
            
            results.append({
                "index": i,
                "original": original,
                "corrected": final_text,
                "explanation": explanation,
                "note": note or "",
                "has_changes": original != final_text
            })
            
            corrected_paragraphs.append(final_text)
        
        return jsonify({
            "success": True,
            "model_used": model,
            "pipeline_used": pipeline,
            "qwen_model_used": qwen_variant,
            "ollama_model_used": ollama_model_name,
            "total_paragraphs": len(paragraphs),
            "results": results,
            "full_corrected": '\n\n'.join(corrected_paragraphs)
        })
        
    except Exception as e:
        import traceback
        return jsonify({
            "success": False,
            "error": str(e),
            "traceback": traceback.format_exc()
        }), 500


@app.route('/api/upload-docx', methods=['POST'])
def upload_docx():
    """
    Upload và đọc nội dung file DOCX
    """
    try:
        if 'file' not in request.files:
            return jsonify({
                "success": False,
                "error": "No file uploaded"
            }), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({
                "success": False,
                "error": "No file selected"
            }), 400
        
        if not file.filename.endswith('.docx'):
            return jsonify({
                "success": False,
                "error": "Only .docx files are supported"
            }), 400
        
        # Đọc file DOCX
        from docx import Document
        import io
        
        doc = Document(io.BytesIO(file.read()))
        
        # Lấy text từ tất cả paragraphs
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        text = '\n'.join(paragraphs)
        
        return jsonify({
            "success": True,
            "filename": file.filename,
            "text": text,
            "paragraph_count": len(paragraphs)
        })
        
    except Exception as e:
        import traceback
        return jsonify({
            "success": False,
            "error": str(e),
            "traceback": traceback.format_exc()
        }), 500


@app.route('/api/download-docx', methods=['POST'])
def download_docx():
    """
    Tạo file DOCX từ văn bản
    """
    try:
        data = request.get_json()
        if not data or 'text' not in data:
            return jsonify({
                "success": False,
                "error": "Missing 'text' field"
            }), 400
        
        text = data['text'].strip()
        filename = data.get('filename', 'corrected_output.docx')
        
        if not text:
            return jsonify({
                "success": False,
                "error": "Text cannot be empty"
            }), 400
        
        # Tạo file DOCX
        from docx import Document
        from flask import send_file
        import io
        
        doc = Document()
        
        # Thêm các đoạn văn
        for para in text.split('\n\n'):
            if para.strip():
                doc.add_paragraph(para.strip())
        
        # Lưu vào memory buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        import traceback
        return jsonify({
            "success": False,
            "error": str(e),
            "traceback": traceback.format_exc()
        }), 500


@app.route('/api/correct-docx', methods=['POST'])
def correct_docx():
    """
    Upload DOCX, sửa lỗi, và trả về DOCX với comments ghi chú thay đổi.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        import io
        
        if 'file' not in request.files:
            return jsonify({
                "success": False,
                "error": "No file uploaded"
            }), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({
                "success": False,
                "error": "No file selected"
            }), 400
        
        if not file.filename.endswith('.docx'):
            return jsonify({
                "success": False,
                "error": "Only .docx files are supported"
            }), 400
        
        # Lấy model và pipeline từ form data
        model = request.form.get('model', DEFAULT_MODEL).lower()
        pipeline = request.form.get('pipeline', DEFAULT_PIPELINE)
        qwen_variant = request.form.get('qwen_model', None)
        
        if model not in AVAILABLE_MODELS:
            model = DEFAULT_MODEL
        if pipeline not in PIPELINE_STRATEGIES:
            pipeline = DEFAULT_PIPELINE
        
        # Đọc file DOCX
        doc = Document(io.BytesIO(file.read()))
        
        # Tạo document mới với nội dung đã sửa
        new_doc = Document()
        changes_log = []
        
        for para_idx, para in enumerate(doc.paragraphs):
            original_text = para.text.strip()
            
            if not original_text:
                new_doc.add_paragraph()
                continue
            
            # Kiểm tra đoạn văn có ý nghĩa để xử lý hay không
            if not is_meaningful_text(original_text):
                # Bỏ qua đoạn không có ý nghĩa, giữ nguyên
                new_doc.add_paragraph(original_text)
                continue
            
            # Sửa lỗi với pipeline
            final_text, explanation = correct_with_pipeline(original_text, model=model, pipeline=pipeline, qwen_variant=qwen_variant)
            
            # Thêm paragraph đã sửa
            new_para = new_doc.add_paragraph(final_text)
            
            # Nếu có thay đổi, ghi chú
            if original_text != final_text:
                changes_log.append({
                    "paragraph": para_idx + 1,
                    "original": original_text,
                    "corrected": final_text,
                    "explanation": explanation
                })
        
        # Thêm phần tổng kết thay đổi ở cuối
        if changes_log:
            new_doc.add_paragraph()
            summary_para = new_doc.add_paragraph()
            summary_run = summary_para.add_run("═══ TỔNG KẾT CÁC THAY ĐỔI ═══")
            summary_run.bold = True
            summary_run.font.size = Pt(14)
            summary_run.font.color.rgb = RGBColor(0, 102, 204)
            
            for change in changes_log:
                new_doc.add_paragraph()
                
                # Tiêu đề đoạn
                title_para = new_doc.add_paragraph()
                title_run = title_para.add_run(f"📍 Đoạn {change['paragraph']}:")
                title_run.bold = True
                
                # Văn bản gốc
                orig_para = new_doc.add_paragraph()
                orig_run = orig_para.add_run("❌ Gốc: ")
                orig_run.font.color.rgb = RGBColor(204, 0, 0)
                orig_para.add_run(change['original'][:200] + "..." if len(change['original']) > 200 else change['original'])
                
                # Văn bản đã sửa
                corr_para = new_doc.add_paragraph()
                corr_run = corr_para.add_run("✅ Sửa: ")
                corr_run.font.color.rgb = RGBColor(0, 153, 0)
                corr_para.add_run(change['corrected'][:200] + "..." if len(change['corrected']) > 200 else change['corrected'])
                
                # Giải thích
                if change['explanation']:
                    exp_para = new_doc.add_paragraph()
                    exp_run = exp_para.add_run("💬 Chú thích: ")
                    exp_run.italic = True
                    exp_para.add_run(change['explanation'])
        
        # Lưu vào buffer
        buffer = io.BytesIO()
        new_doc.save(buffer)
        buffer.seek(0)
        
        # Tạo tên file output
        output_filename = file.filename.replace('.docx', '_corrected.docx')
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name=output_filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        import traceback
        return jsonify({
            "success": False,
            "error": str(e),
            "traceback": traceback.format_exc()
        }), 500


if __name__ == '__main__':
    print("🚀 Starting Vietnamese Text Corrector API...")
    print("📍 API running at: http://localhost:5000")
    print("📖 Endpoints:")
    print("   GET  /api/health - Health check (shows available models & pipelines)")
    print("   POST /api/correct - Correct single text (sync)")
    print("   POST /api/correct-paragraphs - Correct multiple paragraphs (sync)")
    print("   POST /api/submit-job - Submit job to queue (async)")
    print("   GET  /api/job-status/<id> - Get job status/result")
    print("   GET  /api/queue-status - Get queue statistics")
    print("   POST /api/upload-docx - Upload DOCX file")
    print("   POST /api/download-docx - Download as DOCX")
    print("   POST /api/correct-docx - Upload & correct DOCX with comments")
    print("=" * 50)
    print(f"🤖 Available models: {AVAILABLE_MODELS}")
    print(f"🔧 Available pipelines: {PIPELINE_STRATEGIES}")
    print(f"📊 Max queue size: {MAX_QUEUE_SIZE}")
    print("=" * 50)
    
    # Enable threaded mode for concurrent request handling
    app.run(host='0.0.0.0', port=5000, debug=False, threaded=True)
