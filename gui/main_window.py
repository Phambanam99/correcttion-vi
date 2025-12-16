import sys
import os
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QTextEdit, QLabel, QFileDialog, QSplitter,
    QTableWidget, QTableWidgetItem, QHeaderView, QMessageBox,
    QProgressBar, QGroupBox
)
from PyQt5.QtCore import Qt
from PyQt5.QtGui import QFont, QColor
from docx import Document

from gui.log_handler import redirect_stdout_to_gui, restore_stdout
from gui.worker_thread import CorrectionWorker


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.worker = None
        self.changes_list = []  # Lưu trữ các thay đổi
        self.init_ui()
        self.setup_log_redirect()
    
    def init_ui(self):
        self.setWindowTitle("📝 Vietnamese Text Corrector")
        self.setGeometry(100, 100, 1200, 800)
        self.setStyleSheet("""
            QMainWindow {
                background-color: #1e1e2e;
            }
            QLabel {
                color: #cdd6f4;
                font-size: 14px;
                font-weight: bold;
            }
            QTextEdit {
                background-color: #313244;
                color: #cdd6f4;
                border: 2px solid #45475a;
                border-radius: 8px;
                padding: 10px;
                font-size: 13px;
            }
            QPushButton {
                background-color: #89b4fa;
                color: #1e1e2e;
                border: none;
                border-radius: 8px;
                padding: 12px 20px;
                font-size: 14px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #b4befe;
            }
            QPushButton:disabled {
                background-color: #45475a;
                color: #6c7086;
            }
            QTableWidget {
                background-color: #313244;
                color: #cdd6f4;
                border: 2px solid #45475a;
                border-radius: 8px;
                gridline-color: #45475a;
            }
            QTableWidget::item {
                padding: 8px;
            }
            QHeaderView::section {
                background-color: #45475a;
                color: #cdd6f4;
                padding: 8px;
                border: none;
                font-weight: bold;
            }
            QGroupBox {
                color: #cdd6f4;
                font-weight: bold;
                border: 2px solid #45475a;
                border-radius: 8px;
                margin-top: 10px;
                padding-top: 10px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 15px;
                padding: 0 5px;
            }
        """)
        
        # Central widget
        central = QWidget()
        self.setCentralWidget(central)
        main_layout = QVBoxLayout(central)
        main_layout.setSpacing(15)
        main_layout.setContentsMargins(20, 20, 20, 20)
        
        # === BUTTONS PANEL ===
        btn_layout = QHBoxLayout()
        
        self.btn_open = QPushButton("📂 Mở File DOCX")
        self.btn_open.clicked.connect(self.open_file)
        btn_layout.addWidget(self.btn_open)
        
        self.btn_process = QPushButton("▶️ Sửa lỗi")
        self.btn_process.clicked.connect(self.start_processing)
        btn_layout.addWidget(self.btn_process)
        
        self.btn_save = QPushButton("💾 Lưu kết quả")
        self.btn_save.clicked.connect(self.save_file)
        self.btn_save.setEnabled(False)
        btn_layout.addWidget(self.btn_save)
        
        self.btn_clear = QPushButton("🗑️ Xóa")
        self.btn_clear.clicked.connect(self.clear_all)
        self.btn_clear.setStyleSheet("""
            QPushButton {
                background-color: #f38ba8;
            }
            QPushButton:hover {
                background-color: #eba0ac;
            }
        """)
        btn_layout.addWidget(self.btn_clear)
        
        btn_layout.addStretch()
        main_layout.addLayout(btn_layout)
        
        # === MAIN SPLITTER ===
        main_splitter = QSplitter(Qt.Vertical)
        
        # Top section: Input/Output
        top_widget = QWidget()
        top_layout = QHBoxLayout(top_widget)
        top_layout.setContentsMargins(0, 0, 0, 0)
        
        # Input text
        input_group = QGroupBox("📥 VĂN BẢN GỐC")
        input_layout = QVBoxLayout(input_group)
        self.text_input = QTextEdit()
        self.text_input.setPlaceholderText("Dán văn bản cần sửa vào đây hoặc click 'Mở File DOCX'...")
        input_layout.addWidget(self.text_input)
        top_layout.addWidget(input_group)
        
        # Output text
        output_group = QGroupBox("📤 ĐÃ CHỈNH SỬA")
        output_layout = QVBoxLayout(output_group)
        self.text_output = QTextEdit()
        self.text_output.setPlaceholderText("Kết quả sau khi sửa sẽ hiển thị ở đây...")
        self.text_output.setReadOnly(True)
        output_layout.addWidget(self.text_output)
        top_layout.addWidget(output_group)
        
        main_splitter.addWidget(top_widget)
        
        # Middle section: Changes table
        changes_group = QGroupBox("🔄 CÁC THAY ĐỔI")
        changes_layout = QVBoxLayout(changes_group)
        
        self.changes_table = QTableWidget()
        self.changes_table.setColumnCount(3)
        self.changes_table.setHorizontalHeaderLabels(["#", "❌ TRƯỚC", "✅ SAU"])
        self.changes_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Stretch)
        self.changes_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.Stretch)
        self.changes_table.setColumnWidth(0, 50)
        self.changes_table.cellClicked.connect(self.on_change_selected)
        changes_layout.addWidget(self.changes_table)
        
        main_splitter.addWidget(changes_group)
        
        # Explanation section: Chú thích/Giải thích (như comment)
        explain_group = QGroupBox("💬 CHÚ THÍCH / GIẢI THÍCH")
        explain_layout = QVBoxLayout(explain_group)
        
        self.explanation_text = QTextEdit()
        self.explanation_text.setReadOnly(True)
        self.explanation_text.setMaximumHeight(120)
        self.explanation_text.setPlaceholderText("Click vào một dòng trong bảng 'CÁC THAY ĐỔI' để xem giải thích...")
        self.explanation_text.setStyleSheet("""
            QTextEdit {
                background-color: #1e1e2e;
                border: 2px solid #f9e2af;
                border-radius: 8px;
                color: #f9e2af;
                font-size: 13px;
            }
        """)
        explain_layout.addWidget(self.explanation_text)
        
        main_splitter.addWidget(explain_group)
        
        # Bottom section: Log
        log_group = QGroupBox("📋 LOG")
        log_layout = QVBoxLayout(log_group)
        
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        self.log_text.setMaximumHeight(150)
        self.log_text.setStyleSheet("""
            QTextEdit {
                background-color: #11111b;
                font-family: 'Consolas', 'Courier New', monospace;
                font-size: 12px;
            }
        """)
        log_layout.addWidget(self.log_text)
        
        main_splitter.addWidget(log_group)
        
        # Set splitter sizes
        main_splitter.setSizes([350, 150, 100, 120])
        main_layout.addWidget(main_splitter)
    
    def setup_log_redirect(self):
        """Redirect print() vào log panel"""
        self.log_signal = redirect_stdout_to_gui()
        self.log_signal.message.connect(self.append_log)
    
    def append_log(self, message: str):
        """Thêm message vào log panel"""
        self.log_text.append(message.rstrip())
        # Auto scroll to bottom
        scrollbar = self.log_text.verticalScrollBar()
        scrollbar.setValue(scrollbar.maximum())
    
    def open_file(self):
        """Mở file DOCX và đọc nội dung"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Mở File DOCX", "", "Word Documents (*.docx)"
        )
        if file_path:
            try:
                doc = Document(file_path)
                text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
                self.text_input.setPlainText(text)
                self.append_log(f"📂 Đã mở file: {os.path.basename(file_path)}")
            except Exception as e:
                QMessageBox.critical(self, "Lỗi", f"Không thể mở file: {str(e)}")
    
    def start_processing(self):
        """Bắt đầu xử lý sửa lỗi"""
        text = self.text_input.toPlainText().strip()
        if not text:
            QMessageBox.warning(self, "Cảnh báo", "Vui lòng nhập văn bản hoặc mở file DOCX!")
            return
        
        # Disable buttons
        self.btn_process.setEnabled(False)
        self.btn_open.setEnabled(False)
        self.btn_save.setEnabled(False)
        
        # Clear previous results
        self.text_output.clear()
        self.changes_table.setRowCount(0)
        self.changes_list.clear()
        
        # Start worker thread
        self.worker = CorrectionWorker(text)
        self.worker.progress.connect(self.append_log)
        self.worker.paragraph_done.connect(self.on_paragraph_done)
        self.worker.finished.connect(self.on_processing_finished)
        self.worker.error.connect(self.on_processing_error)
        self.worker.start()
    
    def on_paragraph_done(self, index: int, original: str, qwen_result: str, final: str, note: str, explanation: str):
        """Xử lý khi một đoạn văn hoàn thành"""
        if original != final:
            row = self.changes_table.rowCount()
            self.changes_table.insertRow(row)
            self.changes_table.setItem(row, 0, QTableWidgetItem(str(index + 1)))
            self.changes_table.setItem(row, 1, QTableWidgetItem(original[:100] + "..." if len(original) > 100 else original))
            self.changes_table.setItem(row, 2, QTableWidgetItem(final[:100] + "..." if len(final) > 100 else final))
            
            # Color the cells
            self.changes_table.item(row, 1).setBackground(QColor("#45475a"))
            self.changes_table.item(row, 2).setBackground(QColor("#313244"))
            
            self.changes_list.append({
                'index': index,
                'original': original,
                'bartpho': qwen_result,
                'final': final,
                'note': note,
                'explanation': explanation
            })
    
    def on_change_selected(self, row: int, column: int):
        """Hiển thị giải thích khi click vào một thay đổi trong bảng"""
        if row < len(self.changes_list):
            change = self.changes_list[row]
            explanation = change.get('explanation', '')
            note = change.get('note', '')
            
            display_text = ""
            if explanation:
                display_text += f"📝 GIẢI THÍCH:\n{explanation}\n\n"
            if note:
                display_text += f"🔄 GHI CHÚ THAY ĐỔI:\n{note}"
            
            if not display_text:
                display_text = "Không có giải thích cho thay đổi này."
            
            self.explanation_text.setPlainText(display_text)
    
    def on_processing_finished(self, result: str):
        """Xử lý khi hoàn thành toàn bộ"""
        self.text_output.setPlainText(result)
        self.btn_process.setEnabled(True)
        self.btn_open.setEnabled(True)
        self.btn_save.setEnabled(True)
        self.append_log("\n✅ Hoàn thành sửa lỗi!")
    
    def on_processing_error(self, error: str):
        """Xử lý khi có lỗi"""
        self.append_log(error)
        self.btn_process.setEnabled(True)
        self.btn_open.setEnabled(True)
        QMessageBox.critical(self, "Lỗi", error)
    
    def save_file(self):
        """Lưu kết quả ra file DOCX"""
        text = self.text_output.toPlainText().strip()
        if not text:
            QMessageBox.warning(self, "Cảnh báo", "Không có nội dung để lưu!")
            return
        
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Lưu File DOCX", "output_corrected.docx", "Word Documents (*.docx)"
        )
        if file_path:
            try:
                doc = Document()
                for para in text.split('\n\n'):
                    if para.strip():
                        doc.add_paragraph(para.strip())
                doc.save(file_path)
                self.append_log(f"💾 Đã lưu file: {os.path.basename(file_path)}")
                QMessageBox.information(self, "Thành công", f"Đã lưu file: {file_path}")
            except Exception as e:
                QMessageBox.critical(self, "Lỗi", f"Không thể lưu file: {str(e)}")
    
    def clear_all(self):
        """Xóa tất cả nội dung"""
        self.text_input.clear()
        self.text_output.clear()
        self.log_text.clear()
        self.explanation_text.clear()
        self.changes_table.setRowCount(0)
        self.changes_list.clear()
        self.btn_save.setEnabled(False)
    
    def closeEvent(self, event):
        """Cleanup when closing"""
        restore_stdout()
        if self.worker and self.worker.isRunning():
            self.worker.cancel()
            self.worker.wait()
        event.accept()


def run_app():
    """Entry point để khởi động ứng dụng GUI"""
    app = QApplication(sys.argv)
    app.setStyle('Fusion')
    
    window = MainWindow()
    window.show()
    
    sys.exit(app.exec_())


if __name__ == "__main__":
    run_app()
