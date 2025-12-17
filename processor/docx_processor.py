from docx import Document
from llm.qwen_model import correct_text
from protonx_layer.protonx_refine import refine_text
from processor.diff_utils import generate_change_note, is_meaningful_text
from processor.track_comment import add_comment
from config import AUTHOR_NAME

def process_docx(input_path, output_path):
    doc = Document(input_path)
    new_doc = Document()

    total_paragraphs = len([p for p in doc.paragraphs if p.text.strip()])
    print("\n" + "🚀" * 25)
    print(f"📄 Bắt đầu xử lý file: {input_path}")
    print(f"📊 Tổng số đoạn văn cần xử lý: {total_paragraphs}")
    print("🚀" * 25 + "\n")

    para_index = 0
    for para in doc.paragraphs:
        original = para.text.strip()

        if not original:
            new_doc.add_paragraph("")
            continue
        
        # Kiểm tra đoạn văn có ý nghĩa để xử lý hay không
        if not is_meaningful_text(original):
            print(f"⏭️ Bỏ qua đoạn không có ý nghĩa: '{original[:50]}'")
            new_doc.add_paragraph(original)  # Giữ nguyên đoạn gốc
            continue

        para_index += 1
        print("\n" + "🔷" * 25)
        print(f"📝 ĐOẠN VĂN [{para_index}/{total_paragraphs}]")
        print("🔷" * 25)
        print(f"📄 GỐC: {original[:100]}{'...' if len(original) > 100 else ''}")

        # 1️⃣ Qwen sửa ngữ cảnh
        qwen_fixed = correct_text(original)

        # 2️⃣ ProtonX correction cuối
        final_text = refine_text(qwen_fixed)

        # === LOG: Đoạn cần sửa ===
        if original != final_text:
            print("\n" + "⚠️" * 25)
            print("🔄 ĐOẠN CẦN SỬA:")
            print("-" * 50)
            print(f"❌ GỐC    : {original}")
            print(f"✅ ĐÃ SỬA : {final_text}")
            print("⚠️" * 25)
        else:
            print("\n✅ Đoạn văn không cần sửa")

        # 3️⃣ Ghi kết quả
        new_para = new_doc.add_paragraph(final_text)

        # 4️⃣ Track change
        note = generate_change_note(original, final_text)
        if note:
            print(f"📌 Ghi chú thay đổi: {note[:50]}{'...' if len(note) > 50 else ''}")
            add_comment(new_para, note, AUTHOR_NAME)

    print("\n" + "✅" * 25)
    print(f"💾 Đã lưu kết quả vào: {output_path}")
    print("✅" * 25 + "\n")

    new_doc.save(output_path)
